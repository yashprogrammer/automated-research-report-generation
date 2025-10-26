import os
import sys
import re
from datetime import datetime
from typing import Optional
from langgraph.types import Send
from jinja2 import Template

current_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.abspath(os.path.join(current_dir, "../../"))
sys.path.append(project_root)

from langgraph.graph import StateGraph, START, END
from langgraph.checkpoint.memory import MemorySaver
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_community.tools.tavily_search import TavilySearchResults

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from research_and_analyst.schemas.models import (
    Perspectives,
    GenerateAnalystsState,
    ResearchGraphState,
)
from research_and_analyst.utils.model_loader import ModelLoader
from research_and_analyst.workflows.interview_workflow import InterviewGraphBuilder
from research_and_analyst.prompt_lib.prompt_locator import (
    CREATE_ANALYSTS_PROMPT,
    INTRO_CONCLUSION_INSTRUCTIONS,
    REPORT_WRITER_INSTRUCTIONS,
)
from research_and_analyst.logger import GLOBAL_LOGGER
from research_and_analyst.exception.custom_exception import ResearchAnalystException


class AutonomousReportGenerator:
    """
    Handles the end-to-end autonomous report generation workflow using LangGraph.
    """

    def __init__(self, llm):
        self.llm = llm
        self.memory = MemorySaver()
        self.tavily_search = TavilySearchResults(
            tavily_api_key="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"
        )
        self.logger = GLOBAL_LOGGER.bind(module="AutonomousReportGenerator")

    # ----------------------------------------------------------------------
    def create_analyst(self, state: GenerateAnalystsState):
        """Generate analyst personas based on topic and feedback."""
        topic = state["topic"]
        max_analysts = state["max_analysts"]
        human_analyst_feedback = state.get("human_analyst_feedback", "")

        try:
            self.logger.info("Creating analyst personas", topic=topic)
            structured_llm = self.llm.with_structured_output(Perspectives)
            system_prompt = CREATE_ANALYSTS_PROMPT.render(
                topic=topic, max_analysts=max_analysts,
                human_analyst_feedback=human_analyst_feedback,
            )
            analysts = structured_llm.invoke([
                SystemMessage(content=system_prompt),
                HumanMessage(content="Generate the set of analysts."),
            ])
            self.logger.info("Analysts created", count=len(analysts.analysts))
            return {"analysts": analysts.analysts}
        except Exception as e:
            self.logger.error("Error creating analysts", error=str(e))
            raise ResearchAnalystException("Failed to create analysts", e)

    # ----------------------------------------------------------------------
    def human_feedback(self):
        """Pause node for human analyst feedback."""
        try:
            self.logger.info("Awaiting human feedback")
        except Exception as e:
            self.logger.error("Error during feedback stage", error=str(e))
            raise ResearchAnalystException("Human feedback node failed", e)

    # ----------------------------------------------------------------------
    def write_report(self, state: ResearchGraphState):
        """Compile all report sections into unified content."""
        sections = state.get("sections", [])
        topic = state.get("topic", "")

        try:
            if not sections:
                sections = ["No sections generated — please verify interview stage."]
            self.logger.info("Writing report", topic=topic)
            system_prompt = REPORT_WRITER_INSTRUCTIONS.render(topic=topic)
            report = self.llm.invoke([
                SystemMessage(content=system_prompt),
                HumanMessage(content="\n\n".join(sections))
            ])
            self.logger.info("Report written successfully")
            return {"content": report.content}
        except Exception as e:
            self.logger.error("Error writing main report", error=str(e))
            raise ResearchAnalystException("Failed to write main report", e)

    # ----------------------------------------------------------------------
    def write_introduction(self, state: ResearchGraphState):
        """Generate the report introduction."""
        try:
            sections = state["sections"]
            topic = state["topic"]
            formatted_str_sections = "\n\n".join([f"{s}" for s in sections])
            self.logger.info("Generating introduction", topic=topic)
            system_prompt = INTRO_CONCLUSION_INSTRUCTIONS.render(
                topic=topic, formatted_str_sections=formatted_str_sections
            )
            intro = self.llm.invoke([
                SystemMessage(content=system_prompt),
                HumanMessage(content="Write the report introduction")
            ])
            self.logger.info("Introduction generated", length=len(intro.content))
            return {"introduction": intro.content}
        except Exception as e:
            self.logger.error("Error generating introduction", error=str(e))
            raise ResearchAnalystException("Failed to generate introduction", e)

    # ----------------------------------------------------------------------
    def write_conclusion(self, state: ResearchGraphState):
        """Generate the conclusion section."""
        try:
            sections = state["sections"]
            topic = state["topic"]
            formatted_str_sections = "\n\n".join([f"{s}" for s in sections])
            self.logger.info("Generating conclusion", topic=topic)
            system_prompt = INTRO_CONCLUSION_INSTRUCTIONS.render(
                topic=topic, formatted_str_sections=formatted_str_sections
            )
            conclusion = self.llm.invoke([
                SystemMessage(content=system_prompt),
                HumanMessage(content="Write the report conclusion")
            ])
            self.logger.info("Conclusion generated", length=len(conclusion.content))
            return {"conclusion": conclusion.content}
        except Exception as e:
            self.logger.error("Error generating conclusion", error=str(e))
            raise ResearchAnalystException("Failed to generate conclusion", e)

    # ----------------------------------------------------------------------
    def finalize_report(self, state: ResearchGraphState):
        """Assemble introduction, content, and conclusion into final report."""
        try:
            content = state["content"]
            self.logger.info("Finalizing report compilation")
            if content.startswith("## Insights"):
                content = content.strip("## Insights")

            sources = None
            if "## Sources" in content:
                try:
                    content, sources = content.split("\n## Sources\n")
                except Exception:
                    pass

            final_report = (
                state["introduction"] + "\n\n---\n\n" +
                content + "\n\n---\n\n" +
                state["conclusion"]
            )
            if sources:
                final_report += "\n\n## Sources\n" + sources

            self.logger.info("Report finalized")
            return {"final_report": final_report}
        except Exception as e:
            self.logger.error("Error finalizing report", error=str(e))
            raise ResearchAnalystException("Failed to finalize report", e)

    # ----------------------------------------------------------------------
    def save_report(self, final_report: str, topic: str,
                    format: str = "docx"):
        """Save the report as DOCX or PDF, each in its own subfolder."""
        try:
            self.logger.info("Saving report", topic=topic, format=format)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_topic = re.sub(r'[\\/*?:"<>|]', "_", topic)
            base_name = f"{safe_topic.replace(' ', '_')}_{timestamp}"

            # Root folder (always inside project)
            project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), "../../"))
            root_dir = os.path.join(project_root, "generated_report")

            # Create subfolder for this report
            report_folder = os.path.join(root_dir, base_name)
            os.makedirs(report_folder, exist_ok=True)

            # Final file path inside that folder
            file_path = os.path.join(report_folder, f"{base_name}.{format}")

            if format == "docx":
                self._save_as_docx(final_report, file_path)
            elif format == "pdf":
                self._save_as_pdf(final_report, file_path)
            else:
                raise ValueError("Invalid format. Use 'docx' or 'pdf'.")

            self.logger.info("Report saved successfully", path=file_path)
            return file_path

        except Exception as e:
            self.logger.error("Error saving report", error=str(e))
            raise ResearchAnalystException("Failed to save report file", e)

    # ----------------------------------------------------------------------
    def _save_as_docx(self, text: str, file_path: str):
        """Helper: save as DOCX."""
        try:
            doc = Document()
            for line in text.split("\n"):
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                else:
                    doc.add_paragraph(line)
            doc.save(file_path)
        except Exception as e:
            self.logger.error("DOCX save failed", path=file_path, error=str(e))
            raise ResearchAnalystException("Error saving DOCX report", e)

    def _save_as_pdf(self, text: str, file_path: str):
        """Helper: save as PDF."""
        try:
            c = canvas.Canvas(file_path, pagesize=letter)
            _, height = letter
            x, y = 50, height - 50
            for line in text.split("\n"):
                if not line.strip():
                    y -= 15
                    continue
                if y < 50:
                    c.showPage()
                    y = height - 50
                if line.startswith("# "):
                    c.setFont("Helvetica-Bold", 14)
                    line = line[2:]
                elif line.startswith("## "):
                    c.setFont("Helvetica-Bold", 12)
                    line = line[3:]
                else:
                    c.setFont("Helvetica", 10)
                c.drawString(x, y, line.strip())
                y -= 15
            c.save()
        except Exception as e:
            self.logger.error("PDF save failed", path=file_path, error=str(e))
            raise ResearchAnalystException("Error saving PDF report", e)

    # ----------------------------------------------------------------------
    def build_graph(self):
        """Construct the report generation graph."""
        try:
            self.logger.info("Building report generation graph")
            builder = StateGraph(ResearchGraphState)
            interview_graph = InterviewGraphBuilder(self.llm, self.tavily_search).build()

            def initiate_all_interviews(state: ResearchGraphState):
                topic = state.get("topic", "Untitled Topic")
                analysts = state.get("analysts", [])
                if not analysts:
                    self.logger.warning("No analysts found — skipping interviews")
                    return END
                return [
                    Send(
                        "conduct_interview",
                        {
                            "analyst": analyst,
                            "messages": [HumanMessage(content=f"So, let's discuss about {topic}.")],
                            "max_num_turns": 2,
                            "context": [],
                            "interview": "",
                            "sections": [],
                        },
                    )
                    for analyst in analysts
                ]

            builder.add_node("create_analyst", self.create_analyst)
            builder.add_node("human_feedback", self.human_feedback)
            builder.add_node("conduct_interview", interview_graph)
            builder.add_node("write_report", self.write_report)
            builder.add_node("write_introduction", self.write_introduction)
            builder.add_node("write_conclusion", self.write_conclusion)
            builder.add_node("finalize_report", self.finalize_report)

            builder.add_edge(START, "create_analyst")
            builder.add_edge("create_analyst", "human_feedback")
            builder.add_conditional_edges(
                "human_feedback",
                initiate_all_interviews,
                ["conduct_interview", END]
            )
            builder.add_edge("conduct_interview", "write_report")
            builder.add_edge("conduct_interview", "write_introduction")
            builder.add_edge("conduct_interview", "write_conclusion")
            builder.add_edge(["write_report", "write_introduction", "write_conclusion"], "finalize_report")
            builder.add_edge("finalize_report", END)

            graph = builder.compile(interrupt_before=["human_feedback"], checkpointer=self.memory)
            self.logger.info("Report generation graph built successfully")
            return graph
        except Exception as e:
            self.logger.error("Error building report graph", error=str(e))
            raise ResearchAnalystException("Failed to build report generation graph", e)


# ----------------------------------------------------------------------
if __name__ == "__main__":
    try:
        llm = ModelLoader().load_llm()
        reporter = AutonomousReportGenerator(llm)
        graph = reporter.build_graph()

        topic = "Impact of LLMs over the Future of Jobs?"
        thread = {"configurable": {"thread_id": "1"}}
        reporter.logger.info("Starting report generation pipeline", topic=topic)

        for _ in graph.stream({"topic": topic, "max_analysts": 3}, thread, stream_mode="values"):
            pass

        state = graph.get_state(thread)
        feedback = input("\n Enter your feedback or press Enter to continue: ").strip()
        graph.update_state(thread, {"human_analyst_feedback": feedback}, as_node="human_feedback")

        for _ in graph.stream(None, thread, stream_mode="values"):
            pass

        final_state = graph.get_state(thread)
        final_report = final_state.values.get("final_report")

        if final_report:
            reporter.logger.info("Report generated successfully")
            reporter.save_report(final_report, topic, "docx")
            reporter.save_report(final_report, topic, "pdf")
        else:
            reporter.logger.error("No report content generated")

    except Exception as e:
        GLOBAL_LOGGER.error("Fatal error in main execution", error=str(e))
        raise ResearchAnalystException("Autonomous report generation pipeline failed", e)
