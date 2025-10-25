import os
import sys
from datetime import datetime
from typing import Optional
from langgraph.types import Send

current_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.abspath(os.path.join(current_dir, "../../"))
sys.path.append(project_root)

from langgraph.graph import StateGraph, START, END
from langgraph.checkpoint.memory import MemorySaver
from langchain_core.messages import HumanMessage, SystemMessage, AIMessage
from langchain_core.messages import get_buffer_string
from langchain_community.tools.tavily_search import TavilySearchResults

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from research_and_analyst.backend_server.models import (
    Analyst,
    Perspectives,
    GenerateAnalystsState,
    InterviewState,
    ResearchGraphState,
    SearchQuery
    
)

from research_and_analyst.utils.model_loader import ModelLoader
from research_and_analyst.prompt_lib.prompts import *


def build_interview_graph(llm,tavily_search=None):
    
    """Create a LangGraph subgraph that handles interviews."""

    memory = MemorySaver()
    
    def generation_question(state:InterviewState):
        """_summary_

        Args:
            state (InterviewState): _description_
        """
        analyst = state["analyst"]
        messages = state["messages"]
        
        #generate the question
        system_message = ANALYST_ASK_QUESTIONS.format(goals = analyst.persona)
        question = llm.invoke([SystemMessage(content=system_message)]+messages)
        
        #returen the question through state
        return {"messages":[question]}
        
    def search_web(state: InterviewState):
        """_summary_

        Args:
            state (InterviewState): _description_
        """
        structure_llm = llm.with_structured_output(SearchQuery)
        search_query = structure_llm.invoke([GENERATE_SEARCH_QUERY]+state["messages"])
        
        # Search
        search_docs = tavily_search.invoke(search_query.search_query,)
        print("********************")
        print(search_docs)
        # Format
        formatted_search_docs = "\n\n---\n\n".join(
            [
                f'<Document href="{doc["url"]}"/>\n{doc["content"]}\n</Document>'
                for doc in search_docs
            ]
        )

        return {"context": [formatted_search_docs]}
    
    def generate_answer(state: InterviewState):
        """_summary_

        Args:
            state (InterviewState): _description_
        """
            # Get state
        analyst = state["analyst"]
        messages = state["messages"]
        context = state["context"]

        # Answer question
        system_message = GENERATE_ANSWERS.format(goals=analyst.persona, context=context)
        answer = llm.invoke([SystemMessage(content=system_message)]+messages)
                
        # Name the message as coming from the expert
        answer.name = "expert"
        
        # Append it to state
        return {"messages": [answer]}
    
    
    def save_interview(state: InterviewState):
        """_summary_

        Args:
            state (InterviewState): _description_
        """
            # Get messages
        messages = state["messages"]
        
        # Convert interview to a string
        interview = get_buffer_string(messages)
        
        # Save to interviews key
        return {"interview": interview}
    
    def write_section(state: InterviewState):
        """_summary_

        Args:
            state (InterviewState): _description_
        """
         # Get state
        context = state["context"]
        analyst = state["analyst"]
    
        # Write section using either the gathered source docs from interview (context) or the interview itself (interview)
        system_message = WRITE_SECTION.format(focus=analyst.description)
        section = llm.invoke([SystemMessage(content=system_message)]+[HumanMessage(content=f"Use this source to write your section: {context}")]) 
                    
        # Append it to state
        return {"sections": [section.content]}

    builder = StateGraph(InterviewState)
    builder.add_node("ask_question", generation_question)
    builder.add_node("search_web", search_web)
    builder.add_node("generate_answer", generate_answer)
    builder.add_node("save_interview", save_interview)
    builder.add_node("write_section", write_section)

    builder.add_edge(START, "ask_question")
    builder.add_edge("ask_question", "search_web")
    builder.add_edge("search_web", "generate_answer")
    builder.add_edge("generate_answer", "save_interview")
    builder.add_edge("save_interview", "write_section")
    builder.add_edge("write_section", END)

    return builder.compile(checkpointer=memory)


class AutonomousReportGenerator:
    def __init__(self,llm):
        """_summary_
        """
        self.llm = llm
        self.memory = MemorySaver()
        self.tavily_search = TavilySearchResults(tavily_api_key="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")
    
    def create_analyst(self,state:GenerateAnalystsState):
        """_summary_
        """
        topic = state["topic"]
        max_analysts = state["max_analysts"]
        human_analyst_feedback = state.get("human_analyst_feedback","")
        
        structured_llm = llm.with_structured_output(Perspectives)
        
        system_messages = CREATE_ANALYSTS_PROMPT.format(
            topic=topic,
            max_analysts=max_analysts,
            human_analyst_feedback=human_analyst_feedback
            
            )
        analysts = structured_llm.invoke([SystemMessage(content=system_messages)]+ [HumanMessage(content="Generate the set of analysts.")])
        
        # Write the list of analysis to state
        return {"analysts": analysts.analysts}
    
    def human_feedback(self):
        """_summary_
        """
        pass
    
    def write_report(self,state: ResearchGraphState):
        """_summary_
        """
        sections = state.get("sections", [])
        topic = state.get("topic", "")
        system_message = f"You are compiling a unified research report on: {topic}."
        if not sections:
            sections = ["No sections generated — please verify interview stage."]
    
        report = self.llm.invoke([
            SystemMessage(content=system_message),
            HumanMessage(content="\n\n".join(sections))
        ])
        return {"content": report.content}
            
    def write_introduction(self,state:ResearchGraphState):
            # Full set of sections
        sections = state["sections"]
        topic = state["topic"]

        # Concat all sections together
        formatted_str_sections = "\n\n".join([f"{section}" for section in sections])
        
        # Summarize the sections into a final report
        
        instructions = INTRO_CONCLUSION_INSTRUCTIONS.format(topic=topic, formatted_str_sections=formatted_str_sections)    
        intro = llm.invoke([instructions]+[HumanMessage(content=f"Write the report introduction")]) 
        return {"introduction": intro.content}

    
    def write_conclusion(self,state:ResearchGraphState):
        """_summary_
        """
        sections = state["sections"]
        topic = state["topic"]

        # Concat all sections together
        formatted_str_sections = "\n\n".join([f"{section}" for section in sections])
        
        # Summarize the sections into a final report
        
        instructions = INTRO_CONCLUSION_INSTRUCTIONS.format(topic=topic, formatted_str_sections=formatted_str_sections)    
        conclusion = llm.invoke([instructions]+[HumanMessage(content=f"Write the report conclusion")]) 
        return {"conclusion": conclusion.content}
    
    def finalize_report(self,state:ResearchGraphState):
        """_summary_
        """
        content = state["content"]
        if content.startswith("## Insights"):
            content = content.strip("## Insights")
        if "## Sources" in content:
            try:
                content, sources = content.split("\n## Sources\n")
            except Exception as e:
                sources = None
        else:
            sources = None

        final_report = state["introduction"] + "\n\n---\n\n" + content + "\n\n---\n\n" + state["conclusion"]
        if sources is not None:
            final_report += "\n\n## Sources\n" + sources
        return {"final_report": final_report}
        
    def save_report(self,final_report: str, topic: str, format: str = "docx", save_dir: str = None):
        """_summary_
        """
        
        import re
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        # Sanitize topic for Windows file system
        safe_topic = re.sub(r'[\\/*?:"<>|]', "_", topic)
        filename = f"{safe_topic.replace(' ', '_')}_{timestamp}.{format}"
        
        if save_dir is None:
            save_dir =  os.path.join(os.getcwd(),"generated_report")
        os.makedirs(save_dir,exist_ok=True)
        file_path = os.path.join(save_dir,filename)
        
        if format == "docx":
            self._save_as_docx(final_report, file_path)
            
        elif format == "pdf":
            self._save_as_pdf(final_report,file_path)
            
        else:
            raise ValueError("Invalid format. Use 'docx' or 'pdf'.")
        
        print(f"Report Saved: {file_path}")
        return file_path
            
    
    def _save_as_docx(self, text: str, file_path: str):
        doc = Document()
        for line in text.split("\n"):
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            else:
                doc.add_paragraph(line)
        doc.save(file_path)

    def _save_as_pdf(self, text: str, file_path: str):
        c = canvas.Canvas(file_path, pagesize=letter)
        width, height = letter
        x, y = 50, height - 50
        for line in text.split("\n"):
            if not line.strip():
                y -= 15
                continue
            if y < 50:
                c.showPage()
                y = height - 50
            if line.startswith("# "):
                c.setFont("Helvetica-Bold", 14)
                line = line[2:]
            elif line.startswith("## "):
                c.setFont("Helvetica-Bold", 12)
                line = line[3:]
            else:
                c.setFont("Helvetica", 10)
            c.drawString(x, y, line.strip())
            y -= 15
        c.save()
    
    def build_graph(self):
        """_summary_
        """
        
        builder = StateGraph(ResearchGraphState)
        
        interview_graph = build_interview_graph(self.llm, self.tavily_search)
        
        def initiate_all_interviews(state: ResearchGraphState):
            topic = state["topic"]
            analysts = state.get("analysts", [])
            if not analysts:
                print("No analysts found — skipping interviews.")
                return END
            # Create one Send() event per analyst
            return [
                Send(
                    "conduct_interview",
                    {
                        "analyst": analyst,
                        "messages": [HumanMessage(content=f"So, let's discuss about {topic}.")],
                        "max_num_turns": 2,
                        "context": [],
                        "interview": "",
                        "sections": [],
                    },
                )
                for analyst in analysts
            ]
        
        # Add nodes
        builder.add_node("create_analyst", self.create_analyst)
        builder.add_node("human_feedback", self.human_feedback)
        builder.add_node("conduct_interview", interview_graph)
        builder.add_node("write_report", self.write_report)
        builder.add_node("write_introduction", self.write_introduction)
        builder.add_node("write_conclusion", self.write_conclusion)
        builder.add_node("finalize_report", self.finalize_report)

        # Edges
        builder.add_edge(START, "create_analyst")
        builder.add_edge("create_analyst", "human_feedback")

        # Map each analyst → interview graph
        builder.add_conditional_edges("human_feedback", initiate_all_interviews, ["conduct_interview"])

        builder.add_edge("conduct_interview", "write_report")
        builder.add_edge("conduct_interview", "write_introduction")
        builder.add_edge("conduct_interview", "write_conclusion")
        builder.add_edge(
            ["write_report", "write_introduction", "write_conclusion"], "finalize_report"
        )
        builder.add_edge("finalize_report", END)

        return builder.compile(interrupt_before=["human_feedback"], checkpointer=self.memory)
    
if __name__ == "__main__":
        """_summary_
        """
        llm = ModelLoader().load_llm()
        
        reporter = AutonomousReportGenerator(llm)
        
        graph = reporter.build_graph()
        
        topic = "Impact of GenAI over the future Jobs?"
        
        thread = {"configurable": {"thread_id": "1"}}
        
        for _ in graph.stream({"topic": topic, "max_analysts": 3}, thread, stream_mode="values"):
            """_summary_
            """
            pass
        
        state = graph.get_state(thread)
        
        feedback = input("\n Enter your feedback or press Enter to continue as is: ").strip()

        graph.update_state(thread, {"human_analyst_feedback": feedback}, as_node="human_feedback")
        
        for _ in graph.stream(None, thread, stream_mode="values"):pass

        
        final_state = graph.get_state(thread)
        final_report = final_state.values.get("final_report")
        
        if final_report:
            reporter.save_report(final_report, topic, "docx")
            reporter.save_report(final_report, topic, "pdf")
            
        else:
            print("No Report Content Generated")